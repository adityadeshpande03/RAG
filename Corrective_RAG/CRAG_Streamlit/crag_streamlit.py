import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.schema import Document
from typing import List, Dict, Any
from chromadb import PersistentClient
import re
import os
import tempfile
import json
import docx

# Initialize Streamlit page configuration
st.set_page_config(page_title="RAG System", layout="wide")
st.title("RAG System with Groq, LangChain, and Tavily")
st.write("Ask a question and optionally provide files (PDF, TXT, JSON, DOCX) for additional context!")

# Define the RAG prompt template
prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a helpful AI assistant. Use the following context to answer the question. If you're unsure or the context doesn't contain relevant information, say so."),
    ("user", "Context: {context}\n\nQuestion: {question}")
])

class RAGSystem:
    def __init__(self, groq_api_key: str, tavily_api_key: str):
        self.temp_dir = tempfile.mkdtemp()
        
        self.text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=250,
            chunk_overlap=0
        )
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-l6-v2"
        )
        self.llm = ChatGroq(
            api_key=groq_api_key,
            model='llama-3.3-70b-versatile',
            temperature=0
        )
        os.environ["TAVILY_API_KEY"] = tavily_api_key
        self.web_search_tool = TavilySearchResults()
        self.vectorstore = None
        
    def validate_urls(self, urls_input: str) -> List[str]:
        if not urls_input:
            return []
        urls = [url.strip() for url in urls_input.split(',')]
        return [url for url in urls if re.match(r'^(http|https)://', url)]
        
    def load_documents(self, urls: List[str]) -> List[Document]:
        if not urls:
            return []
        try:
            docs = []
            for url in urls:
                loader = WebBaseLoader(url)
                docs.extend(loader.load())
            return self.text_splitter.split_documents(docs)
        except Exception as e:
            st.error(f"Error loading documents: {str(e)}")
            return []
    
    def load_pdf(self, pdf_file) -> List[Document]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(pdf_file.read())
            temp_pdf_path = temp_pdf.name
        loader = PyPDFLoader(temp_pdf_path)
        return self.text_splitter.split_documents(loader.load())

    def load_txt(self, txt_file) -> List[Document]:
        content = txt_file.read().decode("utf-8")
        return [Document(page_content=content)]
    
    def load_json(self, json_file) -> List[Document]:
        content = json.load(json_file)
        text_content = json.dumps(content, indent=2)  # Format JSON for readability
        return [Document(page_content=text_content)]
    
    def load_docx(self, docx_file) -> List[Document]:
        doc = docx.Document(docx_file)
        text_content = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text_content)]
    
    def setup_vectorstore(self, documents: List[Document]) -> bool:
        if not documents:
            return False
        try:
            chroma_client = PersistentClient(path=self.temp_dir)
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                client=chroma_client,
                collection_name='rag-chroma'
            )
            return True
        except Exception as e:
            st.error(f"Error setting up vector store: {str(e)}")
            return False
            
    def web_search(self, question: str) -> List[Document]:
        try:
            results = self.web_search_tool.invoke({"query": question})
            return [Document(page_content=result["content"]) for result in results] if results else []
        except Exception as e:
            st.error(f"Error during web search: {str(e)}")
            return []
    
    def process_query(self, question: str, use_web_search: bool = True) -> Dict[str, Any]:
        try:
            docs = []
            if self.vectorstore:
                docs.extend(self.vectorstore.as_retriever().get_relevant_documents(question))
            if use_web_search:
                docs.extend(self.web_search(question))
            if not docs:
                return {"error": "No relevant information found."}
            context = "\n".join([doc.page_content for doc in docs])
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({"context": context, "question": question})
            return {"response": response, "documents": docs}
        except Exception as e:
            return {"error": f"Error processing query: {str(e)}"}

# Streamlit UI components
with st.sidebar:
    st.header("API Configuration")
    groq_api_key = st.text_input("Enter Groq API Key", type="password")
    tavily_api_key = st.text_input("Enter Tavily API Key", type="password")
    use_web_search = st.checkbox("Enable Web Search", value=True)

question = st.text_input("Enter your question")
urls_input = st.text_area("Enter URLs for additional context (optional, separate by commas)")
uploaded_files = st.file_uploader("Upload files for additional context (PDF, TXT, JSON, DOCX)", type=["pdf", "txt", "json", "docx"], accept_multiple_files=True)

if st.button("Process"):
    if not groq_api_key or not tavily_api_key:
        st.error("Please provide both API keys in the sidebar.")
    elif not question:
        st.error("Please enter a question.")
    else:
        rag_system = RAGSystem(groq_api_key, tavily_api_key)
        valid_urls = rag_system.validate_urls(urls_input)
        url_documents = rag_system.load_documents(valid_urls)
        
        documents = url_documents
        
        for uploaded_file in uploaded_files:
            if uploaded_file.type == "application/pdf":
                documents.extend(rag_system.load_pdf(uploaded_file))
            elif uploaded_file.type == "text/plain":
                documents.extend(rag_system.load_txt(uploaded_file))
            elif uploaded_file.type == "application/json":
                documents.extend(rag_system.load_json(uploaded_file))
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                documents.extend(rag_system.load_docx(uploaded_file))
        
        if documents:
            rag_system.setup_vectorstore(documents)
        result = rag_system.process_query(question, use_web_search)
        if "error" in result:
            st.error(result["error"])
        else:
            st.success("Query processed successfully!")
            st.write("Response:", result["response"])
            with st.expander("View Retrieved Documents"):
                for i, doc in enumerate(result["documents"]):
                    st.write(f"Document {i+1}:")
                    st.write(doc.page_content)
                    st.write("---")
