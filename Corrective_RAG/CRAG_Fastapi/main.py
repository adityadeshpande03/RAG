from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.schema import Document
from chromadb import PersistentClient
import re
import os
import tempfile
import json
import docx
from io import BytesIO
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set API keys
os.environ["GROQ_API_KEY"] = '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'
os.environ["TAVILY_API_KEY"] = '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'

app = FastAPI(title="RAG System API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define request/response models
class QueryRequest(BaseModel):
    question: str
    urls: Optional[List[str]] = []
    use_web_search: bool = True

class QueryResponse(BaseModel):
    response: str
    documents: List[str]

# Initialize the prompt template
prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a helpful AI assistant. Use the following context to answer the question. If you're unsure or the context doesn't contain relevant information, say so."),
    ("user", "Context: {context}\n\nQuestion: {question}")
])

class RAGSystem:
    def __init__(self):
        try:
            self.temp_dir = tempfile.mkdtemp()
            logger.info("Initializing RAG System...")
            
            self.text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=250,
                chunk_overlap=0
            )
            
            logger.info("Loading embeddings model...")
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-l6-v2"
            )
            
            logger.info("Initializing Groq LLM...")
            self.llm = ChatGroq(
                api_key=os.getenv("GROQ_API_KEY"),
                model='llama-3.3-70b-versatile',  # Original model
                temperature=0
            )
            
            logger.info("Initializing Tavily search...")
            self.web_search_tool = TavilySearchResults()
            self.vectorstore = None
            
            logger.info("RAG System initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing RAG System: {str(e)}")
            raise

    def validate_urls(self, urls: List[str]) -> List[str]:
        return [url for url in urls if re.match(r'^(http|https)://', url)]

    def load_documents(self, urls: List[str]) -> List[Document]:
        if not urls:
            return []
        try:
            docs = []
            for url in urls:
                loader = WebBaseLoader(url)
                docs.extend(loader.load())
            return self.text_splitter.split_documents(docs)
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error loading documents: {str(e)}")

    def load_pdf(self, file_content: bytes) -> List[Document]:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(file_content)
                temp_pdf_path = temp_pdf.name
            loader = PyPDFLoader(temp_pdf_path)
            return self.text_splitter.split_documents(loader.load())
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error loading PDF: {str(e)}")

    def load_txt(self, file_content: bytes) -> List[Document]:
        try:
            content = file_content.decode("utf-8")
            return [Document(page_content=content)]
        except Exception as e:
            logger.error(f"Error loading TXT: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error loading TXT: {str(e)}")

    def load_json(self, file_content: bytes) -> List[Document]:
        try:
            content = json.loads(file_content)
            text_content = json.dumps(content, indent=2)
            return [Document(page_content=text_content)]
        except Exception as e:
            logger.error(f"Error loading JSON: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error loading JSON: {str(e)}")

    def load_docx(self, file_content: bytes) -> List[Document]:
        try:
            doc = docx.Document(BytesIO(file_content))
            text_content = "\n".join([para.text for para in doc.paragraphs])
            return [Document(page_content=text_content)]
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error loading DOCX: {str(e)}")

    def setup_vectorstore(self, documents: List[Document]) -> bool:
        if not documents:
            return False
        try:
            chroma_client = PersistentClient(path=self.temp_dir)
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                client=chroma_client,
                collection_name='rag-chroma'
            )
            return True
        except Exception as e:
            logger.error(f"Error setting up vector store: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error setting up vector store: {str(e)}")

    def web_search(self, question: str) -> List[Document]:
        try:
            results = self.web_search_tool.invoke({"query": question})
            return [Document(page_content=result["content"]) for result in results] if results else []
        except Exception as e:
            logger.error(f"Error during web search: {str(e)}")
            return []

    def process_query(self, question: str, use_web_search: bool = True) -> Dict[str, Any]:
        try:
            docs = []
            if use_web_search:
                logger.info("Performing web search...")
                web_docs = self.web_search(question)
                docs.extend(web_docs)
                logger.info(f"Found {len(web_docs)} documents from web search")
            
            if not docs:
                logger.warning("No documents found")
                context = "No additional context found. Answering based on general knowledge."
            else:
                context = "\n".join([doc.page_content for doc in docs])
            
            logger.info("Generating response using LLM...")
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({"context": context, "question": question})
            
            return {
                "response": response,
                "documents": [doc.page_content for doc in docs] if docs else []
            }
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

# API endpoints
@app.post("/api/query", response_model=QueryResponse)
async def query(
    question: str = Form(...),
    use_web_search: bool = Form(True),
    urls: str = Form(""),
    files: List[UploadFile] = File(None)
):
    try:
        logger.info(f"Received query request - Question: {question}")
        rag_system = RAGSystem()
        
        # Process URLs
        url_list = [url.strip() for url in urls.split(',')] if urls else []
        valid_urls = rag_system.validate_urls(url_list)
        documents = rag_system.load_documents(valid_urls)
        
        # Process uploaded files
        if files:
            for file in files:
                content = await file.read()
                if file.filename.endswith('.pdf'):
                    documents.extend(rag_system.load_pdf(content))
                elif file.filename.endswith('.txt'):
                    documents.extend(rag_system.load_txt(content))
                elif file.filename.endswith('.json'):
                    documents.extend(rag_system.load_json(content))
                elif file.filename.endswith('.docx'):
                    documents.extend(rag_system.load_docx(content))

        if documents:
            rag_system.setup_vectorstore(documents)

        result = rag_system.process_query(question, use_web_search)
        return QueryResponse(**result)
    except Exception as e:
        logger.error(f"Error in query endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Health check endpoint
@app.get("/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
